import os
from collections.abc import Generator
from pathlib import Path

import openpyxl
from docx import Document
from loguru import logger


def split_docx(
    file_path: Path, output_path: Path, max_paragraphs=10
) -> Generator[Path]:
    """
    Split .docx document by 'max_paragraphs'.
    """
    input_file = Path(file_path)
    if not input_file.exists():
        msg = f"File {file_path} not found."
        raise FileNotFoundError(msg)
    if not str(file_path).endswith(".docx"):
        msg = f"Файл {file_path} должен быть формата .docx."
        raise ValueError(msg)
    output_path.mkdir(parents=True, exist_ok=True)
    doc = Document(str(file_path))
    paragraphs = doc.paragraphs
    part_num = 1
    part_text = []
    result_files = []
    for i, paragraph in enumerate(paragraphs):
        if not paragraph.text.strip():
            continue
        part_text.append(paragraph.text)
        if (i + 1) % max_paragraphs == 0 or i == len(paragraphs) - 1:
            part_doc = Document()
            for text in part_text:
                part_doc.add_paragraph(text)
            part_output_path = f"{str(output_path)}/part_{part_num}.docx"
            part_doc.save(part_output_path)
            logger.info(
                f"Сохранён файл: {part_output_path} с {len(part_text)} параграфами."
            )
            result_files.append(part_output_path)
            part_text = []
            part_num += 1
            yield Path(part_output_path)


def split_docx_by_size_generator(
    file_path: Path, output_path: Path, max_size_mb: int
) -> Generator[Path]:
    """
    split .docx document by size
    """
    max_size_bytes = max_size_mb * 1024 * 1024
    input_file = Path(file_path)
    if not input_file.exists():
        msg = f"Файл {file_path} не найден."
        raise FileNotFoundError(msg)
    if not str(file_path).endswith(".docx"):
        msg = f"Файл {file_path} должен быть формата .docx."
        raise ValueError(msg)
    output_path.mkdir(parents=True, exist_ok=True)
    doc = Document(str(file_path))
    part_num = 1
    part_doc = Document()
    for paragraph in doc.paragraphs:
        part_doc.add_paragraph(paragraph.text, style=paragraph.style)
        temp_path = output_path / "temp_part.docx"
        part_doc.save(str(temp_path))
        if os.path.getsize(temp_path) >= max_size_bytes:
            part_file_path = output_path / f"part_{part_num}.docx"
            part_doc.save(str(part_file_path))
            yield Path(part_file_path)
            part_doc = Document()
            part_num += 1
    if part_doc.paragraphs:
        part_file_path = output_path / f"part_{part_num}.docx"
        part_doc.save(str(part_file_path))
        yield part_file_path


def split_xlsx(file_path: Path, output_dir: Path) -> Generator[Path]:
    """
    Split .xlsx document by pages.
    """
    workbook = openpyxl.load_workbook(file_path)
    output_dir.mkdir(parents=True, exist_ok=True)
    for sheet_name in workbook.sheetnames:
        worksheet = workbook[sheet_name]
        temp_file = output_dir / f"{file_path.stem}_{sheet_name}.xlsx"
        temp_workbook = openpyxl.Workbook()
        temp_sheet = temp_workbook.active
        temp_sheet.title = sheet_name
        for row in worksheet.iter_rows(values_only=True):
            temp_sheet.append(row)
        temp_workbook.save(temp_file)
        yield temp_file
